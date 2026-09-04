"""Researcher — chat with your engineering codes & standards.

Run:  python run.py   then open http://localhost:8600
"""
import json
import os
import re
import shutil
import subprocess
import sys
import threading
import time
import uuid
from concurrent.futures import ThreadPoolExecutor
from pathlib import Path
import logging

logging.getLogger("pypdf").setLevel(logging.ERROR)

from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import FileResponse, StreamingResponse
from pydantic import BaseModel

from . import store, ingest, rag
from .config import load_config, save_config, ensure_dirs, FILES_DIR, BASE_DIR, DATA_DIR
from .embeddings import embed_in_batches, EmbeddingError, list_models
from .providers import stream_chat, test_connection, ProviderError

app = FastAPI(title="Researcher")
STATIC_DIR = Path(__file__).resolve().parent / "static"

ensure_dirs()
store.get_conn()


# ---------------- indexing pipeline ----------------

def _index_document(doc_id: int):
    cfg = load_config()
    doc = store.get_document(doc_id)
    if not doc:
        return
    try:
        store.set_status(doc_id, "parsing")
        blocks, pages = ingest.parse_file(doc["stored_path"], doc["format"])

        def _no_text_msg():
            try:
                has_ocr = ingest._get_ocr() is not None
            except Exception:
                has_ocr = False
            if doc["format"] == "pdf" and not has_ocr:
                return ("This looks like a scanned PDF, and OCR is not available in this Python "
                        "environment. Use Settings -> 'Check & update dependencies', restart the app, "
                        "then re-index this document.")
            return "No extractable text found in this document."

        if not blocks:
            store.set_status(doc_id, "error", _no_text_msg())
            return
        rcfg = cfg["retrieval"]
        chunks = ingest.chunk_blocks(blocks, rcfg.get("chunk_size", 1500), rcfg.get("chunk_overlap", 200))
        if not chunks:
            store.set_status(doc_id, "error", _no_text_msg())
            return
        store.clear_chunks(doc_id)
        chunk_ids = store.add_chunks(doc_id, chunks)
        store.set_doc_meta(doc_id, pages=pages, chunk_count=len(chunks))

        ecfg = dict(cfg["embeddings"])
        if not ecfg.get("api_key") and cfg.get("chat", {}).get("api_key"):
            ecfg["api_key"] = cfg["chat"]["api_key"]
        if ecfg.get("provider") != "local" and cfg.get("chat", {}).get("provider") == "gemini":
            ecfg["provider"] = "gemini"

        if ecfg.get("enabled", True):

            store.set_status(doc_id, "embedding")
            try:
                embs = embed_in_batches(
                    [c["text"] for c in chunks], ecfg,
                    batch_size=ecfg.get("batch_size", 32),
                )
                store.add_vectors(list(zip(chunk_ids, embs)))
                store.set_status(doc_id, "ready")
            except EmbeddingError as e:
                # keyword search still works
                store.set_status(doc_id, "ready_keyword_only", str(e))
        else:
            store.set_status(doc_id, "ready_keyword_only")
    except Exception as e:
        import traceback
        err_msg = f"{type(e).__name__}: {e}"
        print(f"Indexing error for doc {doc_id}:\n{traceback.format_exc()}")
        store.set_status(doc_id, "error", err_msg)



_index_executor = ThreadPoolExecutor(max_workers=2, thread_name_prefix="doc-indexer")

def _index_async(doc_id: int):
    _index_executor.submit(_index_document, doc_id)


# ---------------- watch folders: auto-index new/changed files ----------------

def _scan_watch_folders(sync: bool = True) -> dict:
    """One scan pass over the configured watch folders. Returns counts."""
    added = updated = 0
    wcfg = load_config().get("watch", {})
    for folder in wcfg.get("folders", []):
        p = Path(str(folder).strip())
        if not p.is_dir():
            continue
        for f in p.rglob("*"):
            try:
                if not f.is_file():
                    continue
                if f.name.startswith((".", "~$")) or f.name in ("Thumbs.db", "desktop.ini"):
                    continue
                fmt = ingest.detect_format(f.name)
                if fmt == "unknown":
                    continue
                mtime = f.stat().st_mtime
                existing = store.get_doc_by_path(str(f))
                if existing is None:
                    doc_id = store.add_document(f.name, str(f), fmt)
                    store.set_doc_mtime(doc_id, mtime)
                    (_index_document if sync else _index_async)(doc_id)
                    added += 1
                elif existing.get("file_mtime") is None or mtime > (existing.get("file_mtime") or 0) + 1:
                    store.set_doc_mtime(existing["id"], mtime)
                    (_index_document if sync else _index_async)(existing["id"])
                    updated += 1
            except Exception:
                continue
        return {"added": added, "updated": updated}


def _watch_loop():
    while True:
        try:
            wcfg = load_config().get("watch", {})
            if wcfg.get("auto", False) and wcfg.get("folders"):
                _scan_watch_folders(sync=True)
            interval = max(30, int(wcfg.get("interval", 60)))
        except Exception:
            interval = 60
        time.sleep(interval)


threading.Thread(target=_watch_loop, daemon=True).start()


# ---------------- routes: UI ----------------

@app.get("/")
def index():
    # no-store so the browser never serves a stale cached UI after an update
    return FileResponse(STATIC_DIR / "index.html",
                        headers={"Cache-Control": "no-store, max-age=0"})


# ---------------- routes: documents ----------------

@app.get("/api/documents")
def list_documents():
    return {"documents": store.list_documents(), "stats": store.stats()}


@app.post("/api/upload")
async def upload(files: list[UploadFile] = File(...)):
    results = []
    for f in files:
        fmt = ingest.detect_format(f.filename)
        if fmt == "unknown":
            results.append({"filename": f.filename, "error": "Unsupported file type"})
            continue
        safe_name = f"{uuid.uuid4().hex[:8]}_{Path(f.filename).name}"
        dest = FILES_DIR / safe_name
        with dest.open("wb") as out:
            shutil.copyfileobj(f.file, out)
        doc_id = store.add_document(Path(f.filename).name, str(dest), fmt)
        _index_async(doc_id)
        results.append({"filename": f.filename, "doc_id": doc_id})
    return {"results": results}


@app.delete("/api/documents/{doc_id}")
def delete_document(doc_id: int):
    doc = store.get_document(doc_id)
    if not doc:
        raise HTTPException(404)
    try:
        sp = Path(doc["stored_path"])
        if FILES_DIR.resolve() in sp.resolve().parents:
            sp.unlink(missing_ok=True)   # only delete files the app itself stored
    except OSError:
        pass
    store.delete_document(doc_id)
    return {"ok": True}


@app.delete("/api/documents")
def clear_all_documents():
    docs = store.list_documents()
    for doc in docs:
        try:
            sp = Path(doc["stored_path"])
            if FILES_DIR.resolve() in sp.resolve().parents:
                sp.unlink(missing_ok=True)
        except OSError:
            pass
    store.clear_all_documents()
    return {"ok": True}



@app.post("/api/documents/{doc_id}/reindex")
def reindex(doc_id: int):
    if not store.get_document(doc_id):
        raise HTTPException(404)
    store.set_status(doc_id, "pending")
    _index_async(doc_id)
    return {"ok": True}


@app.get("/api/documents/{doc_id}/file")
def get_file(doc_id: int):
    doc = store.get_document(doc_id)
    if not doc or not Path(doc["stored_path"]).exists():
        raise HTTPException(404)
    if doc["format"] == "pdf":
        # inline -> browser renders its PDF viewer and honors #page=N
        return FileResponse(doc["stored_path"], media_type="application/pdf",
                            filename=doc["filename"], content_disposition_type="inline")
    return FileResponse(doc["stored_path"], filename=doc["filename"])


@app.post("/api/watch/scan")
def watch_scan():
    """Manually scan the watch folders now (new docs index in the background)."""
    return {"ok": True, **_scan_watch_folders(sync=False)}


@app.get("/api/browse")
def browse(path: str = ""):
    """List sub-folders for the folder picker."""
    if not path:
        if os.name == "nt":
            import string
            drives = [f"{d}:\\" for d in string.ascii_uppercase if Path(f"{d}:\\").exists()]
            return {"path": "", "parent": None, "dirs": drives}
        path = "/"
    p = Path(path)
    if not p.is_dir():
        raise HTTPException(404, "Not a folder")
    dirs = []
    try:
        for x in sorted(p.iterdir(), key=lambda v: v.name.lower()):
            try:
                if x.is_dir() and not x.name.startswith((".", "$")):
                    dirs.append(str(x))
            except OSError:
                continue
            if len(dirs) >= 400:
                break
    except PermissionError:
        raise HTTPException(403, "Permission denied")
    parent = str(p.parent)
    if parent == str(p):
        parent = "" if os.name == "nt" else None
    return {"path": str(p), "parent": parent, "dirs": dirs}


# ---------------- routes: settings / providers ----------------

@app.get("/api/settings")
def get_settings():
    return load_config()


class SettingsBody(BaseModel):
    chat: dict | None = None
    embeddings: dict | None = None
    retrieval: dict | None = None
    watch: dict | None = None


@app.post("/api/settings")
def set_settings(body: SettingsBody):
    cfg = save_config({k: v for k, v in body.model_dump().items() if v is not None})
    return cfg


@app.get("/api/test/chat")
def test_chat():
    return test_connection(load_config()["chat"])


@app.get("/api/test/embeddings")
def test_embeddings():
    cfg = load_config()
    ecfg = dict(cfg.get("embeddings", {}))
    if not ecfg.get("api_key") and cfg.get("chat", {}).get("api_key"):
        ecfg["api_key"] = cfg["chat"]["api_key"]
    if ecfg.get("provider") != "local" and cfg.get("chat", {}).get("provider") == "gemini":
        ecfg["provider"] = "gemini"

    try:
        from .embeddings import embed_texts
        emb = embed_texts(["connection test"], ecfg)
        prov = ecfg.get("provider", "local")
        return {"ok": True, "dim": len(emb[0]), "provider": prov}
    except Exception as e:
        return {"ok": False, "error": str(e)}



class ModelsProbeBody(BaseModel):
    provider: str = "openai-compatible"
    base_url: str = ""
    api_key: str = ""


@app.post("/api/models")
def models_probe(body: ModelsProbeBody):
    """Probe a provider with explicit credentials (used by the Settings test button)."""
    return test_connection(body.model_dump())


@app.get("/api/models")
def models(which: str = "chat"):
    cfg = load_config()[which if which in ("chat", "embeddings") else "chat"]
    try:
        if cfg.get("provider") == "gemini":
            return test_connection(cfg)
        return {"ok": True, "models": list_models(cfg.get("base_url", ""), cfg.get("api_key", ""))}
    except Exception as e:
        return {"ok": False, "error": str(e), "models": []}



@app.post("/api/update-deps")
def update_deps():
    """Install/upgrade Python dependencies from requirements.txt, streaming pip's output."""
    def event(obj: dict) -> str:
        return f"data: {json.dumps(obj, ensure_ascii=False)}\n\n"

    def generate():
        req = BASE_DIR / "requirements.txt"
        if not req.exists():
            yield event({"type": "done", "ok": False, "text": "requirements.txt not found"})
            return
        yield event({"type": "line", "text": f"$ {Path(sys.executable).name} -m pip install -r requirements.txt"})
        try:
            proc = subprocess.Popen(
                [sys.executable, "-m", "pip", "install", "--upgrade", "-r", str(req)],
                stdout=subprocess.PIPE, stderr=subprocess.STDOUT,
                text=True, bufsize=1, cwd=str(BASE_DIR),
            )
            for line in proc.stdout:
                line = line.rstrip()
                if line:
                    yield event({"type": "line", "text": line[:300]})
            code = proc.wait()
            if code != 0:
                yield event({"type": "done", "ok": False, "text": f"pip exited with code {code} - see log above"})
                return
            # optional OCR extras - failure is not fatal
            ocr_note = ""
            ocr_req = BASE_DIR / "requirements-ocr.txt"
            if ocr_req.exists():
                yield event({"type": "line", "text": "$ pip install -r requirements-ocr.txt  (optional OCR extras)"})
                proc2 = subprocess.Popen(
                    [sys.executable, "-m", "pip", "install", "--upgrade", "-r", str(ocr_req)],
                    stdout=subprocess.PIPE, stderr=subprocess.STDOUT,
                    text=True, bufsize=1, cwd=str(BASE_DIR),
                )
                for line in proc2.stdout:
                    line = line.rstrip()
                    if line:
                        yield event({"type": "line", "text": line[:300]})
                if proc2.wait() != 0:
                    ocr_note = (" OCR extras failed to install (Python version may be too new) - "
                                "the app still works, scanned pages are skipped.")
            yield event({"type": "done", "ok": True,
                         "text": "Dependencies are up to date. Restart the app to load new packages." + ocr_note})
        except Exception as e:
            yield event({"type": "done", "ok": False, "text": str(e)})

    return StreamingResponse(generate(), media_type="text/event-stream",
                             headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})


# ---------------- routes: collections ----------------

class CollectionBody(BaseModel):
    name: str | None = None


@app.post("/api/documents/{doc_id}/collection")
def set_doc_collection(doc_id: int, body: CollectionBody):
    if not store.get_document(doc_id):
        raise HTTPException(404)
    store.set_collection(doc_id, (body.name or "").strip() or None)
    return {"ok": True}


@app.get("/api/collections")
def collections():
    return {"collections": store.list_collections()}


# ---------------- routes: exports ----------------

def _clean_md(s: str) -> str:
    s = re.sub(r"\*\*([^*]+)\*\*", r"\1", s)
    s = re.sub(r"\*([^*]+)\*", r"\1", s)
    s = re.sub(r"`([^`]+)`", r"\1", s)
    return s.strip()


def _md_to_docx(content: str, sources: list, path: str):
    from docx import Document as DocxDocument
    d = DocxDocument()
    d.add_heading("Researcher answer", level=1)
    lines = content.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue
        if line.lstrip().startswith("#"):
            stripped = line.lstrip()
            lvl = min(len(stripped) - len(stripped.lstrip("#")), 4)
            d.add_heading(_clean_md(stripped.lstrip("# ")), level=max(2, lvl))
        elif line.lstrip().startswith(("- ", "* ")):
            d.add_paragraph(_clean_md(line.lstrip()[2:]), style="List Bullet")
        elif re.match(r"^\s*\d+\.\s", line):
            d.add_paragraph(_clean_md(re.sub(r"^\s*\d+\.\s*", "", line)), style="List Number")
        elif "|" in line and line.strip().startswith("|"):
            rows = []
            while i < len(lines) and "|" in lines[i]:
                if not re.match(r"^\s*\|?[\s:\-|]+\|?\s*$", lines[i]):
                    rows.append([_clean_md(cell) for cell in lines[i].strip().strip("|").split("|")])
                i += 1
            if rows:
                ncol = max(len(r) for r in rows)
                t = d.add_table(rows=len(rows), cols=ncol)
                t.style = "Table Grid"
                for ri, row in enumerate(rows):
                    for ci, cell in enumerate(row):
                        t.cell(ri, ci).text = cell
            continue
        else:
            d.add_paragraph(_clean_md(line))
        i += 1
    if sources:
        d.add_heading("Sources", level=2)
        for s in sources:
            loc = f", {s.get('location')}" if s.get("location") else ""
            d.add_paragraph(f"[{s.get('n')}] {s.get('filename', '')}{loc}", style="List Bullet")
    d.save(path)


class ExportAnswerBody(BaseModel):
    content: str
    sources: list[dict] | None = None


@app.post("/api/export/answer")
def export_answer(body: ExportAnswerBody):
    out_dir = DATA_DIR / "exports"
    out_dir.mkdir(parents=True, exist_ok=True)
    path = out_dir / f"answer-{int(time.time())}.docx"
    _md_to_docx(body.content, body.sources or [], str(path))
    return FileResponse(str(path), filename="answer.docx",
                        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


class ExportTableBody(BaseModel):
    text: str


@app.post("/api/export/table")
def export_table(body: ExportTableBody):
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    for line in body.text.splitlines():
        if "|" not in line:
            continue
        if re.match(r"^\s*\|?[\s:\-|]+\|?\s*$", line):
            continue
        ws.append([c.strip() for c in line.strip().strip("|").split("|")])
    if ws.max_row < 1:
        raise HTTPException(400, "No table rows found")
    out_dir = DATA_DIR / "exports"
    out_dir.mkdir(parents=True, exist_ok=True)
    path = out_dir / f"table-{int(time.time())}.xlsx"
    wb.save(str(path))
    return FileResponse(str(path), filename="table.xlsx",
                        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")


# ---------------- routes: chat history ----------------

class ChatCreateBody(BaseModel):
    title: str = "New chat"


class ChatRenameBody(BaseModel):
    title: str


@app.get("/api/chats")
def chats_list():
    return {"chats": store.list_chats()}


@app.post("/api/chats")
def chats_create(body: ChatCreateBody):
    return {"id": store.create_chat(body.title)}


@app.get("/api/chats/{chat_id}")
def chats_get(chat_id: int):
    return {"messages": store.get_chat_messages(chat_id)}


@app.post("/api/chats/{chat_id}/rename")
def chats_rename(chat_id: int, body: ChatRenameBody):
    store.rename_chat(chat_id, body.title)
    return {"ok": True}


@app.delete("/api/chats/{chat_id}")
def chats_delete(chat_id: int):
    store.delete_chat(chat_id)
    return {"ok": True}


# ---------------- routes: search & chat ----------------

@app.get("/api/search")
def search(q: str, top_k: int = 10):
    cfg = load_config()
    results = rag.retrieve(q, cfg)[:top_k]
    return {"results": results}


class ChatBody(BaseModel):
    messages: list[dict]                 # [{role, content}]
    doc_ids: list[int] | None = None     # optional: restrict to selected documents
    use_library: bool = True
    chat_id: int | None = None           # optional: persist this conversation
    compare: bool = False                # comparison mode across selected documents


@app.post("/api/chat")
def chat(body: ChatBody):
    cfg = load_config()
    question = ""
    for m in reversed(body.messages):
        if m.get("role") == "user":
            question = m.get("content", "")
            break

    if body.chat_id and question:
        store.add_chat_message(body.chat_id, "user", question)

    def event(obj: dict) -> str:
        return f"data: {json.dumps(obj, ensure_ascii=False)}\n\n"

    def generate():
        sources = []
        if body.use_library and question:
            try:
                if body.compare and body.doc_ids and len(body.doc_ids) >= 2:
                    # comparison: retrieve separately per document so each is represented
                    per = max(3, cfg["retrieval"].get("top_k", 10) // len(body.doc_ids))
                    sources = []
                    for did in body.doc_ids:
                        sources += rag.retrieve(question, cfg, doc_ids=[did])[:per]
                else:
                    sources = rag.retrieve(question, cfg, doc_ids=body.doc_ids, history=body.messages)
                sources = rag.fit_sources(sources, cfg, history=body.messages)
            except Exception as e:
                yield event({"type": "warning", "text": f"Retrieval failed: {e}"})
        src_payload = [
            {"n": i + 1, "filename": s["filename"], "location": s.get("location"),
             "doc_id": s["doc_id"], "text": s["text"][:600]}
            for i, s in enumerate(sources)
        ]
        yield event({"type": "sources", "sources": src_payload})
        msgs = rag.build_messages(body.messages, sources, compare=body.compare) if body.use_library else body.messages
        answer_parts = []
        try:
            n_deltas = 0
            for delta in stream_chat(msgs, cfg["chat"]):
                n_deltas += 1
                answer_parts.append(delta)
                yield event({"type": "delta", "text": delta})
            if n_deltas == 0:
                yield event({"type": "error", "text":
                    "The model returned an empty response. Most common cause: its context window is "
                    "too small for the retrieved sources - increase the context length in LM Studio "
                    "AND set the same value in Settings -> 'Model context length'. "
                    "Alternatively lower top-k in Settings."})
            yield event({"type": "done"})
        except ProviderError as e:
            yield event({"type": "error", "text": str(e)})
        finally:
            # persist the answer (also saves partial answers when the user hits Stop)
            if body.chat_id:
                ans = re.sub(r"<think>[\s\S]*?(</think>|$)", "", "".join(answer_parts)).strip()
                if ans:
                    store.add_chat_message(body.chat_id, "assistant", ans, sources=src_payload)

    return StreamingResponse(generate(), media_type="text/event-stream",
                             headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})
