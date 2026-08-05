#!/usr/bin/env python3
"""
Generate Knowledge Base JSON for any local directory or Google Drive folder.
Usage:
  python scripts/build_drive_index.py "C:/path/to/my/drive/folder"
  OR in Google Colab:
  python scripts/build_drive_index.py "/content/drive/MyDrive/MyEngineeringCodes"
"""

import os
import sys
import json
import re
import math
import urllib.request
import urllib.parse
import urllib.error
from pathlib import Path

def extract_pdf_pages(file_path):
    pages = []
    try:
        import pypdf
        reader = pypdf.PdfReader(file_path)
        for idx, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            pages.append({"page": idx + 1, "text": text.strip()})
        if pages:
            return pages
    except Exception as e:
        pass

    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            for idx, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                pages.append({"page": idx + 1, "text": text.strip()})
        if pages:
            return pages
    except Exception as e:
        pass

    return pages

def extract_docx_pages(file_path):
    full_text = []
    try:
        import docx
        doc = docx.Document(file_path)
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                r_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if r_text:
                    full_text.append(r_text)
    except Exception as e:
        print(f"DOCX extraction failed for {file_path}: {e}")

    text = "\n".join(full_text)
    if not text.strip():
        try:
            with open(file_path, "rb") as f:
                content = f.read()
                ascii_strings = re.findall(rb'[\x20-\x7e\t\r\n]{4,}', content)
                text = "\n".join(b.decode("utf-8", errors="ignore").strip() for b in ascii_strings if len(b) > 4 and not b.startswith(b"<w:"))
        except Exception:
            text = ""

    chunks_text = [text[i:i+1000] for i in range(0, len(text), 1000) if text[i:i+1000].strip()]
    return [{"page": idx + 1, "text": c} for idx, c in enumerate(chunks_text)]

def extract_text_pages(file_path):
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = f.read()
        chunks_text = [text[i:i+1000] for i in range(0, len(text), 1000)]
        return [{"page": idx + 1, "text": c} for idx, c in enumerate(chunks_text)]
    except Exception as e:
        return []

def tokenize(text):
    return re.findall(r'\b\w+\b', text.lower())

def chunk_text(pages, filename, max_chars=800, overlap=100):
    chunks = []
    chunk_id = 0
    for page_info in pages:
        page_num = page_info["page"]
        text = page_info["text"]
        if not text:
            continue
        start = 0
        while start < len(text):
            end = start + max_chars
            chunk_str = text[start:end].strip()
            if chunk_str:
                clause_match = re.search(r'(?:(?:para|section|clause|article|part)\s+[\d\.]+|[\d]+\.[\d]+(?:\.[\d]+)?)', chunk_str, re.IGNORECASE)
                clause = clause_match.group(0) if clause_match else f"Page {page_num}"
                chunks.append({
                    "id": f"{filename}_{chunk_id}",
                    "file": filename,
                    "page": page_num,
                    "clause": clause,
                    "text": chunk_str,
                    "tokens": tokenize(chunk_str)
                })
                chunk_id += 1
            start += (max_chars - overlap)
    return chunks

def build_index_for_folder(folder_path):
    target_dir = Path(folder_path)
    if not target_dir.exists():
        print(f"Error: Directory {folder_path} does not exist.")
        sys.exit(1)

    supported_exts = {".pdf", ".docx", ".xlsx", ".csv", ".txt", ".md"}
    all_files = [f for f in target_dir.glob("*") if f.suffix.lower() in supported_exts]
    print(f"Found {len(all_files)} document(s) in {folder_path}")

    documents_meta = []
    all_chunks = []

    for file_path in all_files:
        filename = file_path.name
        ext = file_path.suffix.lower()
        print(f"Indexing: {filename}...")

        if ext == ".pdf":
            pages = extract_pdf_pages(file_path)
        elif ext == ".docx":
            pages = extract_docx_pages(file_path)
        else:
            pages = extract_text_pages(file_path)

        doc_chunks = chunk_text(pages, filename)
        documents_meta.append({
            "filename": filename,
            "page_count": len(pages),
            "chunk_count": len(doc_chunks)
        })
        all_chunks.extend(doc_chunks)

    kb_data = {
        "documents": documents_meta,
        "chunks": all_chunks,
        "generated_at": str(target_dir.stat().st_mtime)
    }

    output_file = target_dir / "knowledge_base.json"
    with open(output_file, "w", encoding="utf-8") as f:
        json.dump(kb_data, f, indent=2)

    print(f"✅ Generated {output_file} successfully! ({len(all_chunks)} chunks from {len(documents_meta)} files)")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python scripts/build_drive_index.py <PATH_TO_FOLDER>")
        sys.exit(1)
    build_index_for_folder(sys.argv[1])
